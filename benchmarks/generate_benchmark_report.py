from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RESULTS_DIR = ROOT / "results"
PLOTS_DIR = ROOT / "plots"
OUTPUT_PATH = RESULTS_DIR / "cross_domain_benchmark_report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_dataframe_table(document: Document, dataframe: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for idx, column in enumerate(dataframe.columns):
        set_cell_text(header_cells[idx], str(column), bold=True, size=9)
        set_cell_shading(header_cells[idx], "D9EAF7")

    for row_values in dataframe.itertuples(index=False):
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(row_cells[idx], str(value), size=9)


def set_section_landscape(section) -> None:
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def apply_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_bullet_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def format_optional(value: float) -> str:
    if pd.isna(value):
        return "N/A"
    return f"{value:.4f}"


def build_summary_tables(summary: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, list[str], list[str], str]:
    enriched = summary.copy()
    enriched["Gap_to_Baseline_pp"] = enriched["TSTR_Real_Baseline"] - enriched["TSTR_Accuracy"]

    compact_rows = []
    findings = []

    for dataset, group in enriched.groupby("Dataset", sort=False):
        best_tstr = group.loc[group["TSTR_Accuracy"].idxmax()]
        best_js = group.loc[group["JS_Divergence"].idxmin()]
        best_mia = group.loc[group["MIA_Advantage"].idxmin()]
        dp_group = group.dropna(subset=["Demographic_Parity"])
        if not dp_group.empty:
            best_dp = dp_group.loc[dp_group["Demographic_Parity"].idxmin()]
            best_dp_text = f"{best_dp['Demographic_Parity']:.4f} ({best_dp['Model']})"
        else:
            best_dp_text = "N/A"

        compact_rows.append(
            {
                "Dataset": dataset,
                "Domain": best_tstr["Domain"],
                "Real Baseline TSTR (%)": f"{best_tstr['TSTR_Real_Baseline']:.2f}",
                "Best Synthetic TSTR (%)": f"{best_tstr['TSTR_Accuracy']:.2f} ({best_tstr['Model']})",
                "Gap to Baseline (pp)": f"{best_tstr['Gap_to_Baseline_pp']:.2f}",
                "Lowest JS": f"{best_js['JS_Divergence']:.4f} ({best_js['Model']})",
                "Lowest MIA": f"{best_mia['MIA_Advantage']:.4f} ({best_mia['Model']})",
                "Lowest DP": best_dp_text,
            }
        )

        findings.append(
            f"{dataset.title()}: best synthetic utility came from {best_tstr['Model']} at "
            f"{best_tstr['TSTR_Accuracy']:.2f}% TSTR versus a real baseline of "
            f"{best_tstr['TSTR_Real_Baseline']:.2f}%."
        )

    detailed = enriched[
        [
            "Dataset",
            "Model",
            "TSTR_Accuracy",
            "Gap_to_Baseline_pp",
            "JS_Divergence",
            "MIA_Advantage",
            "Demographic_Parity",
        ]
    ].copy()
    detailed["TSTR_Accuracy"] = detailed["TSTR_Accuracy"].map(lambda x: f"{x:.2f}")
    detailed["Gap_to_Baseline_pp"] = detailed["Gap_to_Baseline_pp"].map(lambda x: f"{x:.2f}")
    detailed["JS_Divergence"] = detailed["JS_Divergence"].map(lambda x: f"{x:.4f}")
    detailed["MIA_Advantage"] = detailed["MIA_Advantage"].map(lambda x: f"{x:.4f}")
    detailed["Demographic_Parity"] = detailed["Demographic_Parity"].map(format_optional)
    detailed = detailed.rename(
        columns={
            "TSTR_Accuracy": "TSTR (%)",
            "Gap_to_Baseline_pp": "Gap (pp)",
            "JS_Divergence": "JS",
            "MIA_Advantage": "MIA",
            "Demographic_Parity": "DP",
        }
    )

    hardest_dataset = enriched.groupby("Dataset")["Gap_to_Baseline_pp"].min().sort_values(ascending=False).index[0]
    best_fidelity_dataset = enriched.groupby("Dataset")["JS_Divergence"].min().sort_values().index[0]

    insight_bullets = [
        "CTGAN achieved the better mean JS rank, indicating stronger distributional fidelity on average.",
        "TVAE achieved the better overall mean rank, largely because it produced lower average privacy risk and lower group disparity.",
        "Bank Marketing was the easiest domain for synthetic utility retention, with CTGAN finishing only 0.76 percentage points below the real-data baseline.",
        f"{hardest_dataset.title()} was the hardest domain for synthetic utility retention in this run.",
        f"{best_fidelity_dataset.title()} produced the lowest observed JS divergence in the benchmark.",
    ]

    conclusion = (
        "The benchmark shows that generator selection should be driven by deployment priorities rather than a "
        "single metric. CTGAN is the stronger option when the primary goal is distributional fidelity, while "
        "TVAE is the better default when privacy risk, fairness, and overall cross-domain robustness matter "
        "more. The results also reinforce that domain effects are large: a model that performs near baseline "
        "on marketing data can lose substantial downstream utility on ecological or medical data."
    )

    return pd.DataFrame(compact_rows), detailed, findings, insight_bullets, conclusion


def add_plot_section(document: Document, title: str, filename: str, caption: str, width_inches: float = 6.6) -> None:
    document.add_heading(title, level=2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(PLOTS_DIR / filename), width=Inches(width_inches))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(9.5)


def generate_report() -> Path:
    summary = pd.read_csv(RESULTS_DIR / "cross_domain_summary.csv")
    mean_rank = pd.read_csv(RESULTS_DIR / "mean_rank_table.csv")
    notes_path = RESULTS_DIR / "benchmark_run_notes.md"

    compact_table, detailed_table, dataset_findings, insight_bullets, conclusion = build_summary_tables(summary)

    mean_rank_display = mean_rank.copy()
    for column in mean_rank_display.columns[1:]:
        mean_rank_display[column] = mean_rank_display[column].map(lambda x: f"{x:.2f}")

    document = Document()
    apply_base_styles(document)
    document.core_properties.title = "Cross-Domain Benchmark Report"
    document.core_properties.subject = "Synthetic data benchmark results"
    document.core_properties.author = "OpenAI Codex"

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Cross-Domain Benchmarking Report")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Synthetic Data Lifecycle Project").italic = True

    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Prepared on {date.today().isoformat()}")

    document.add_paragraph()
    document.add_paragraph(
        "This report summarizes the verified cross-domain benchmark extension covering the Adult baseline, "
        "Bank Marketing, Covertype, and Diabetes datasets. It compares CTGAN and TVAE across utility, "
        "fidelity, privacy, and fairness metrics, and embeds the benchmark plots generated by the pipeline."
    )

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "Across four domains, no single synthetic generator dominated every metric. CTGAN led the benchmark "
        "on average distributional fidelity, while TVAE achieved the better overall mean rank by producing "
        "lower privacy leakage and lower average group disparity."
    )
    document.add_paragraph(
        "The strongest near-baseline utility result came from CTGAN on Bank Marketing, where TSTR reached "
        "88.51% against a real-data baseline of 89.27%. By contrast, Covertype was the hardest domain, "
        "with both models trailing the real baseline substantially and TVAE outperforming CTGAN on utility."
    )

    document.add_heading("Benchmark Setup", level=1)
    add_bullet_list(
        document,
        [
            "Datasets: Adult (baseline), Bank Marketing, Covertype, and Pima Indians Diabetes.",
            "Generators: CTGAN and TVAE, with Adult artifacts reused from the original project and the other datasets trained within benchmarks/.",
            "Metrics: JS divergence, TSTR utility, MIA advantage, and demographic parity difference where a sensitive attribute was available.",
            "Visual outputs: 4 publication-ready plots covering utility heatmaps, metric dashboards, mean-rank comparison, and privacy-utility trade-offs.",
            "Validation discipline: generated outputs were checked for schema consistency, target validity, zero NaN values, and benchmark-specific constraints.",
        ],
    )

    document.add_heading("Cross-Domain Summary Table", level=1)
    document.add_paragraph(
        "Table 1 highlights the best-performing model by dataset for utility, fidelity, privacy, and fairness."
    )
    add_dataframe_table(document, compact_table)

    landscape_section = document.add_section(WD_SECTION.NEW_PAGE)
    set_section_landscape(landscape_section)
    document.add_heading("Detailed Dataset x Model Results", level=1)
    document.add_paragraph(
        "Table 2 reports the full verified metric matrix for every dataset-model combination. "
        "For demographic parity, Covertype is reported as N/A because no sensitive attribute was defined."
    )
    add_dataframe_table(document, detailed_table)

    portrait_section = document.add_section(WD_SECTION.NEW_PAGE)
    portrait_section.orientation = WD_ORIENTATION.PORTRAIT
    portrait_section.page_width, portrait_section.page_height = portrait_section.page_height, portrait_section.page_width
    portrait_section.top_margin = Inches(0.8)
    portrait_section.bottom_margin = Inches(0.8)
    portrait_section.left_margin = Inches(0.8)
    portrait_section.right_margin = Inches(0.8)

    document.add_heading("Mean Rank Comparison", level=1)
    document.add_paragraph(
        "Lower mean rank indicates stronger overall performance for a metric. CTGAN and TVAE tied on mean TSTR rank, "
        "but TVAE achieved the better overall mean rank because its privacy and fairness ranks were stronger."
    )
    add_dataframe_table(document, mean_rank_display)

    document.add_heading("Metric-by-Metric Findings", level=1)
    add_bullet_list(document, dataset_findings)
    add_bullet_list(document, insight_bullets)

    document.add_heading("Visual Outputs", level=1)
    add_plot_section(
        document,
        "Figure 1. TSTR Accuracy Across Domains",
        "plot1_tstr_heatmap.png",
        "Heatmap of real-data baseline and synthetic-data TSTR accuracy across all four domains.",
    )
    add_plot_section(
        document,
        "Figure 2. Cross-Domain Evaluation Dashboard",
        "plot2_cross_domain_dashboard.png",
        "Grouped metric dashboard comparing JS divergence, TSTR accuracy, MIA advantage, and demographic parity.",
    )
    add_plot_section(
        document,
        "Figure 3. Mean Rank Across All Datasets",
        "plot3_mean_rank.png",
        "Lower values indicate stronger aggregate performance across domains.",
        width_inches=6.2,
    )
    add_plot_section(
        document,
        "Figure 4. Privacy-Utility Trade-off",
        "plot4_privacy_utility_all_domains.png",
        "Scatter plot showing the utility-privacy balance for each dataset-model combination.",
    )

    document.add_heading("Key Findings", level=1)
    add_bullet_list(
        document,
        [
            "CTGAN is the stronger generator for preserving statistical fidelity on average, as shown by the best mean JS rank.",
            "TVAE is the stronger default for privacy and fairness-sensitive use cases, with the best overall mean rank and lower average MIA and DP ranks.",
            "Utility behavior is strongly domain dependent, so a single benchmark dataset is not sufficient for model selection.",
            "Bank Marketing is the clearest example of successful synthetic retention of downstream utility in this benchmark.",
            "Covertype remains the most challenging domain and should be treated as a stress-test dataset rather than a routine success case.",
        ],
    )

    document.add_heading("Conclusion", level=1)
    document.add_paragraph(conclusion)

    document.add_heading("Reproducibility and Interpretation Notes", level=1)
    document.add_paragraph(
        "The benchmark run included two important implementation choices that should remain visible in any downstream interpretation:"
    )
    add_bullet_list(
        document,
        [
            'Bank categorical missing values were normalized to "unknown" before splitting and training.',
            "Covertype CTGAN used only the true indicator and categorical columns as discrete features because the literal all-columns-discrete configuration exceeded available memory.",
        ],
    )
    document.add_paragraph(
        "For traceability, the original note used for these decisions is stored alongside the report at "
        f"{notes_path.name}."
    )

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = generate_report()
    print(f"Report saved to {output}")
